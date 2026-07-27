"""
Pleadly v2 — 智能分块引擎
三种分块策略: 简历结构化 | 项目文档语义 | 证书整文档
"""

import re
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field


# ═══════════════════════════════════════════════════════════
# DATA MODEL
# ═══════════════════════════════════════════════════════════

@dataclass
class Chunk:
    """一个分块。"""
    content: str
    metadata: Dict = field(default_factory=dict)
    chunk_id: str = ""


# ═══════════════════════════════════════════════════════════
# SECTION MARKERS — 简历章节识别
# ═══════════════════════════════════════════════════════════

SECTION_PATTERNS = [
    ("education",      r"(?:教育(?:背景|经历|程度)?|学历|Education|EDUCATION)"),
    ("experience",     r"(?:工作(?:经历|经验)?|实习(?:经历|经验)?|Experience|Work\s*Experience|EMPLOYMENT)"),
    ("projects",       r"(?:项目(?:经历|经验|作品|展示)?|Projects|Portfolio|PROJECTS)"),
    ("skills",         r"(?:专业)?技能(?:特长|证书)?|技术栈|Skills?|Technical\s*Skills?|SKILLS"),
    ("certifications", r"(?:证书|资格(?:证书)?|Certifications?|CERTIFICATIONS?)"),
    ("summary",        r"(?:个人(?:总结|简介|介绍|评价)?|自我(?:评价|介绍|描述)?|求职意向|Summary|Profile|SUMMARY|PROFILE|Objective)"),
    ("awards",         r"(?:获奖(?:经历|情况)?|荣誉|Awards?|Honors?|AWARDS)"),
    ("languages",      r"(?:语言(?:能力)?|Languages?|LANGUAGES?)"),
    ("contact",        r"(?:联系(?:方式)?|Contact|CONTACT)"),
    ("publications",   r"(?:论文|发表|出版|Publications?|PUBLICATIONS?)"),
]


def _detect_section(line: str) -> Optional[str]:
    """检测一行是否是章节标题，返回 section name 或 None。"""
    line = line.strip().rstrip('：:')
    if not line or len(line) > 30:
        return None
    for section_name, pattern in SECTION_PATTERNS:
        if re.search(pattern, line, re.IGNORECASE):
            return section_name
    return None


# ═══════════════════════════════════════════════════════════
# CHUNKING STRATEGIES
# ═══════════════════════════════════════════════════════════

def resume_chunker(
    text: str,
    source_file: str = "unknown",
    resume_version: str = "default",
    min_section_length: int = 30,
) -> List[Chunk]:
    """
    简历结构化分块 — 按章节标题切分。
    每章 = 一个 chunk，带 section 标签。
    """
    lines = text.split('\n')
    chunks: List[Chunk] = []
    current_section = "header"
    current_lines: List[str] = []

    for line in lines:
        section = _detect_section(line)
        if section:
            # Flush previous section
            content = '\n'.join(current_lines).strip()
            if len(content) >= min_section_length:
                chunks.append(Chunk(
                    content=content,
                    metadata={
                        "doc_type": "resume",
                        "section": current_section,
                        "source_file": source_file,
                        "resume_version": resume_version,
                    }
                ))
            current_section = section
            current_lines = [line]
        else:
            current_lines.append(line)

    # Flush last section
    content = '\n'.join(current_lines).strip()
    if len(content) >= min_section_length:
        chunks.append(Chunk(
            content=content,
            metadata={
                "doc_type": "resume",
                "section": current_section,
                "source_file": source_file,
                "resume_version": resume_version,
            }
        ))

    # Assign chunk IDs
    for i, chunk in enumerate(chunks):
        chunk.chunk_id = f"resume_{source_file}_{chunk.metadata['section']}_{i}"

    return chunks


def semantic_chunker(
    text: str,
    doc_type: str = "project",
    source_file: str = "unknown",
    project_name: str = "",
    chunk_size: int = 600,
    overlap: int = 120,
) -> List[Chunk]:
    """
    语义分块 — 按段落 + 重叠窗口切分。
    用于项目文档、长文本描述、作品集等。
    """
    # 先按段落切
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]

    chunks: List[Chunk] = []
    current_chunk: List[str] = []
    current_length = 0

    for para in paragraphs:
        para_len = len(para)

        if current_length + para_len > chunk_size and current_chunk:
            # Flush current chunk
            content = '\n\n'.join(current_chunk)
            chunks.append(content)
            # Overlap: keep last ~overlap chars worth of paragraphs
            overlap_text = content[-overlap:] if len(content) > overlap else content
            current_chunk = [overlap_text.strip()] if overlap_text.strip() else []
            current_length = len(overlap_text) if overlap_text else 0

        current_chunk.append(para)
        current_length += para_len

    # Flush last chunk
    if current_chunk:
        content = '\n\n'.join(current_chunk)
        if content.strip():
            chunks.append(content)

    # Build Chunk objects
    total = len(chunks)
    result = []
    for i, content in enumerate(chunks):
        meta = {
            "doc_type": doc_type,
            "source_file": source_file,
            "chunk_index": i,
            "total_chunks": total,
        }
        if project_name:
            meta["project_name"] = project_name
        result.append(Chunk(
            content=content,
            metadata=meta,
            chunk_id=f"{doc_type}_{source_file}_chunk{i}"
        ))

    return result


def identity_chunker(
    text: str,
    doc_type: str = "other",
    source_file: str = "unknown",
    label: str = "",
) -> List[Chunk]:
    """
    整文档分块 — 一个文档 = 一个 chunk。
    用于证书、链接、短文本等不需要拆分的文档。
    """
    if not text.strip():
        return []

    meta = {
        "doc_type": doc_type,
        "source_file": source_file,
    }
    if label:
        meta["label"] = label

    return [Chunk(
        content=text.strip(),
        metadata=meta,
        chunk_id=f"{doc_type}_{source_file}_full"
    )]


# ═══════════════════════════════════════════════════════════
# AUTO-DETECTION: 根据文本特征自动选分块策略
# ═══════════════════════════════════════════════════════════

def detect_doc_type(text: str, filename: str = "") -> str:
    """根据文本内容和文件名自动检测文档类型。"""
    # 文件名中包含 resume/cv → 简历
    name_lower = filename.lower()
    if any(kw in name_lower for kw in ['resume', '简历', 'cv', 'curriculum']):
        return "resume"

    # 检测简历的章节结构
    section_count = 0
    for line in text.split('\n')[:60]:  # 只看前60行
        if _detect_section(line):
            section_count += 1
    if section_count >= 3:
        return "resume"

    # 文件名中包含 project/作品/项目 → 项目文档
    if any(kw in name_lower for kw in ['project', '项目', '作品', 'portfolio', 'readme']):
        return "project"

    # 文件名中包含 cert/证书 → 证书
    if any(kw in name_lower for kw in ['cert', '证书', 'license', '证书']):
        return "certificate"

    # 短文本 → identity
    if len(text) < 800:
        return "other"

    # 默认 → 语义分块
    return "project"


def auto_chunk(
    text: str,
    source_file: str = "unknown",
    label: str = "",
    chunk_size: int = 600,
    overlap: int = 120,
) -> List[Chunk]:
    """
    自动检测文档类型 → 选择对应分块策略 → 分块。
    外部唯一入口。
    """
    doc_type = detect_doc_type(text, source_file)

    if doc_type == "resume":
        return resume_chunker(text, source_file=source_file, resume_version=label or "default")

    elif doc_type == "project":
        return semantic_chunker(
            text, doc_type="project", source_file=source_file,
            project_name=label or source_file,
            chunk_size=chunk_size, overlap=overlap
        )

    elif doc_type == "certificate":
        return identity_chunker(text, doc_type="certificate", source_file=source_file, label=label)

    else:  # "other"
        return identity_chunker(text, doc_type="other", source_file=source_file, label=label)


# ═══════════════════════════════════════════════════════════
# UTILITY: 从不同格式的文件读取文本
# ═══════════════════════════════════════════════════════════

def read_text_from_file(file_path: str) -> Tuple[str, str]:
    """
    读取文件内容。返回 (文本内容, 文件名)。
    支持: .txt, .md, .docx, .pdf
    """
    import os
    filename = os.path.basename(file_path)

    if not os.path.exists(file_path):
        return "", filename

    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read(), filename
    elif ext == '.md':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read(), filename
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(' | '.join(cells))
            return '\n'.join(paragraphs), filename
        except Exception as e:
            return f"[docx解析失败: {e}]", filename
    elif ext == '.pdf':
        try:
            from pypdf2 import PdfReader
            reader = PdfReader(file_path)
            pages = [page.extract_text() or '' for page in reader.pages]
            return '\n'.join(pages), filename
        except Exception as e:
            return f"[pdf解析失败: {e}]", filename
    else:
        return f"[不支持的文件格式: {ext}]", filename


# ═══════════════════════════════════════════════════════════
# QUERY ROUTER — 判断查询该搜哪些 section
# ═══════════════════════════════════════════════════════════

QUERY_ROUTES: Dict[str, Dict] = {
    "skills": {
        "keywords": ["技能", "技术栈", "会什么", "python", "编程", "开发", "语言", "框架",
                     "skill", "tech", "language", "framework", "tool"],
        "sections": ["skills"],
        "doc_types": ["resume"],
    },
    "experience": {
        "keywords": ["经历", "实习", "工作", "做过", "经验", "担任", "负责",
                     "experience", "work", "intern", "job"],
        "sections": ["experience"],
        "doc_types": ["resume"],
    },
    "education": {
        "keywords": ["学历", "学校", "专业", "毕业", "教育", "本科", "硕士",
                     "education", "school", "university", "degree", "major"],
        "sections": ["education"],
        "doc_types": ["resume"],
    },
    "projects": {
        "keywords": ["项目", "作品", "产品", "开发", "构建", "实现", "交付",
                     "project", "portfolio", "build", "product", "deliver"],
        "sections": ["projects"],
        "doc_types": ["resume", "project"],
    },
    "achievements": {
        "keywords": ["成果", "数据", "结果", "指标", "star", "成就",
                     "achievement", "result", "impact", "metric"],
        "sections": ["experience", "projects", "awards"],
        "doc_types": ["resume", "project"],
    },
    "certificates": {
        "keywords": ["证书", "资格", "认证", "certificate", "license", "certification"],
        "sections": ["certifications"],
        "doc_types": ["resume", "certificate"],
    },
}


class QueryRouter:
    """查询路由器 — 根据查询内容决定检索范围和过滤条件。"""

    def route(self, query: str) -> Dict[str, List[str]]:
        """
        分析查询意图，返回 ChromaDB where 过滤条件。

        Returns:
            {"section": ["skills", "projects"], "doc_type": ["resume", "project"]}
            空 dict = 不需要过滤，全局检索
        """
        query_lower = query.lower()
        matched_routes = []

        for route_name, route_config in QUERY_ROUTES.items():
            for kw in route_config["keywords"]:
                if kw in query_lower:
                    matched_routes.append(route_config)
                    break

        if not matched_routes:
            return {}  # 无匹配，全局检索

        # 合并所有匹配路由的限制
        sections = set()
        doc_types = set()
        for route in matched_routes:
            sections.update(route.get("sections", []))
            doc_types.update(route.get("doc_types", []))

        result = {}
        if sections:
            result["section"] = list(sections)
        if doc_types:
            result["doc_type"] = list(doc_types)

        return result
